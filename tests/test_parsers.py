"""Testes para o módulo de parsers (extração de texto de arquivos)."""

from __future__ import annotations

import csv
import tempfile
from pathlib import Path

import pytest

from legal_anki.parsers import ParseError, parse_file


class TestParseTxt:
    """Testes para leitura de arquivos .txt."""

    def test_parse_txt_basic(self, tmp_path):
        """Lê arquivo .txt corretamente."""
        f = tmp_path / "input.txt"
        f.write_text("Art. 5º da CF/88 garante direitos fundamentais.", encoding="utf-8")

        result = parse_file(f)
        assert "Art. 5º" in result

    def test_parse_txt_utf8_special_chars(self, tmp_path):
        """Preserva caracteres especiais jurídicos (§, º, ª)."""
        f = tmp_path / "input.txt"
        f.write_text("§ 1º do art. 5º — alínea 'a'", encoding="utf-8")

        result = parse_file(f)
        assert "§" in result
        assert "º" in result

    def test_parse_txt_empty_raises_error(self, tmp_path):
        """Arquivo vazio levanta ParseError."""
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")

        with pytest.raises(ParseError, match="vazio"):
            parse_file(f)


class TestParsePdf:
    """Testes para leitura de arquivos .pdf."""

    def test_parse_pdf_with_text(self, tmp_path):
        """Extrai texto de PDF gerado programaticamente."""
        import pymupdf

        pdf_path = tmp_path / "test.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Art. 5º, LXIII, CF/88 - Direito ao silêncio.")
        doc.save(str(pdf_path))
        doc.close()

        result = parse_file(pdf_path)
        assert "Art. 5º" in result
        assert "silêncio" in result

    def test_parse_pdf_multipage(self, tmp_path):
        """Extrai texto de PDF com múltiplas páginas."""
        import pymupdf

        pdf_path = tmp_path / "multi.pdf"
        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Página {i + 1}: conteúdo jurídico relevante.")
        doc.save(str(pdf_path))
        doc.close()

        result = parse_file(pdf_path)
        assert "Página 1" in result
        assert "Página 3" in result

    def test_parse_pdf_empty_raises_error(self, tmp_path):
        """PDF sem texto levanta ParseError."""
        import pymupdf

        pdf_path = tmp_path / "empty.pdf"
        doc = pymupdf.open()
        doc.new_page()  # Página em branco
        doc.save(str(pdf_path))
        doc.close()

        with pytest.raises(ParseError, match="não contém texto"):
            parse_file(pdf_path)


class TestParseDocx:
    """Testes para leitura de arquivos .docx."""

    def test_parse_docx_basic(self, tmp_path):
        """Extrai texto de DOCX simples."""
        import docx

        docx_path = tmp_path / "test.docx"
        doc = docx.Document()
        doc.add_paragraph("Art. 5º da CF/88 garante direitos fundamentais.")
        doc.add_paragraph("Súmula Vinculante 11 - STF.")
        doc.save(str(docx_path))

        result = parse_file(docx_path)
        assert "Art. 5º" in result
        assert "Súmula Vinculante 11" in result

    def test_parse_docx_empty_raises_error(self, tmp_path):
        """DOCX sem parágrafos com texto levanta ParseError."""
        import docx

        docx_path = tmp_path / "empty.docx"
        doc = docx.Document()
        doc.save(str(docx_path))

        with pytest.raises(ParseError, match="vazio"):
            parse_file(docx_path)


class TestParseCsv:
    """Testes para leitura de arquivos .csv."""

    def test_parse_csv_semicolon(self, tmp_path):
        """Lê CSV com separador ponto-e-vírgula (padrão BR)."""
        csv_path = tmp_path / "test.csv"
        csv_path.write_text(
            "topico;conteudo\ndireitos_fundamentais;Art. 5º garante a liberdade\n",
            encoding="utf-8",
        )

        result = parse_file(csv_path)
        assert "Art. 5º" in result

    def test_parse_csv_comma(self, tmp_path):
        """Lê CSV com separador vírgula."""
        csv_path = tmp_path / "test.csv"
        csv_path.write_text(
            "topic,content\nfundamentals,Art. 5 guarantees freedom\n",
            encoding="utf-8",
        )

        result = parse_file(csv_path)
        assert "Art. 5" in result

    def test_parse_csv_no_header(self, tmp_path):
        """Lê CSV sem cabeçalho (dados começam na primeira linha)."""
        csv_path = tmp_path / "test.csv"
        csv_path.write_text(
            "Art. 5º, LXIII, CF/88;Direito ao silêncio\n",
            encoding="utf-8",
        )

        result = parse_file(csv_path)
        assert "Art. 5º" in result

    def test_parse_csv_empty_raises_error(self, tmp_path):
        """CSV vazio levanta ParseError."""
        csv_path = tmp_path / "empty.csv"
        csv_path.write_text("", encoding="utf-8")

        with pytest.raises(ParseError, match="vazio"):
            parse_file(csv_path)


class TestParseFileGeneral:
    """Testes gerais para parse_file."""

    def test_unsupported_extension(self, tmp_path):
        """Extensão não suportada levanta ParseError."""
        f = tmp_path / "test.xlsx"
        f.write_text("data", encoding="utf-8")

        with pytest.raises(ParseError, match="não suportado"):
            parse_file(f)

    def test_file_not_found(self):
        """Arquivo inexistente levanta FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            parse_file(Path("/tmp/nao_existe_12345.txt"))

    def test_whitespace_only_file(self, tmp_path):
        """Arquivo com apenas espaços em branco levanta ParseError."""
        f = tmp_path / "spaces.txt"
        f.write_text("   \n\n  \t  ", encoding="utf-8")

        with pytest.raises(ParseError, match="vazio"):
            parse_file(f)
